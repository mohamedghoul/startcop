"""
Document Parser Module
Handles PDF, DOCX, XLSX, CSV, and image files with OCR
Production-ready with robust error handling
"""

import pdfplumber
from docx import Document
import logging
from typing import Optional, Dict, List, Union
import re
import os
import csv
import io
from PIL import Image
import pytesseract
import pandas as pd
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

class DocumentParser:
    """
    Universal document parser supporting multiple formats
    PDF, DOCX, XLSX, CSV, and images with OCR
    """
    
    def __init__(self, ocr_enabled: bool = True):
        self.ocr_enabled = ocr_enabled
        self.supported_formats = ['.pdf', '.docx', '.xlsx', '.csv', '.jpg', '.jpeg', '.png', '.tiff', '.bmp']
        
        # Configure OCR
        if ocr_enabled:
            try:
                # Test if tesseract is available
                pytesseract.get_tesseract_version()
                logger.info("OCR enabled with tesseract")
            except Exception as e:
                logger.warning(f"OCR not available: {e}")
                self.ocr_enabled = False
    
    def parse_document(self, file_path: str) -> str:
        """
        Main parsing method - routes to appropriate parser based on file type
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {ext}. Supported: {self.supported_formats}")
        
        logger.info(f"Parsing document: {os.path.basename(file_path)}")
        
        # Route to appropriate parser
        parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.xlsx': self._parse_excel,
            '.csv': self._parse_csv,
            '.jpg': self._parse_image,
            '.jpeg': self._parse_image,
            '.png': self._parse_image,
            '.tiff': self._parse_image,
            '.bmp': self._parse_image
        }
        
        return parsers[ext](file_path)
    
    def _parse_pdf(self, file_path: str) -> str:
        """PDF → OCR first (images / no selectable text)"""
        try:
            import pytesseract, pdfplumber
            from PIL import Image

            full_text = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 1. try selectable text first (fast)
                    txt = page.extract_text()
                    if txt and txt.strip():
                        full_text.append(txt.strip())
                        continue

                    # 2. no text → OCR the whole page image
                    im = page.to_image()
                    pil_image = im.original      # PIL image
                    ocr_txt = pytesseract.image_to_string(pil_image, lang='eng')
                    if ocr_txt.strip():
                        full_text.append(ocr_txt.strip())

            return "\n\n".join(full_text)
        except Exception as e:
            logger.warning(f"PDF OCR failed {file_path}: {e}")
            return ""
    

    
    def _parse_pdf_stream(self, file_path: str) -> str:
        "O(n) time, O(1) memory"
        import pdfplumber, gc
        for page in pdfplumber.open(file_path).pages:
            txt = page.extract_text() or ""
            yield self._clean_text(txt)
            gc.collect()          # free page objects immediately


    def _parse_docx(self, file_path: str) -> str:
        """Parse Word document"""
        text_parts = []
        doc = Document(file_path)
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extract from tables
        for table_num, table in enumerate(doc.tables, 1):
            table_text = self._docx_table_to_text(table)
            if table_text:
                text_parts.append(f"Table {table_num}: {table_text}")
        
        return "\n\n".join(text_parts)
    
    def _parse_excel(self, file_path: str) -> str:
        """Parse Excel spreadsheet with multiple sheets"""
        text_parts = []
        
        try:
            # Try with pandas first (faster)
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                if not df.empty:
                    sheet_text = self._dataframe_to_text(df, sheet_name)
                    if sheet_text:
                        text_parts.append(f"Sheet '{sheet_name}':\n{sheet_text}")
        
        except Exception as e:
            logger.warning(f"Pandas Excel parsing failed, trying openpyxl: {e}")
            
            # Fallback to openpyxl
            try:
                wb = load_workbook(file_path, data_only=True)
                
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    sheet_text = self._excel_worksheet_to_text(ws, sheet_name)
                    if sheet_text:
                        text_parts.append(f"Sheet '{sheet_name}':\n{sheet_text}")
                        
            except Exception as e2:
                logger.error(f"Excel parsing failed completely: {e2}")
                raise Exception(f"Failed to parse Excel file: {str(e2)}")
        
        return "\n\n".join(text_parts)
    
    def _parse_csv(self, file_path: str) -> str:
        """Parse CSV file"""
        text_parts = []
        
        with open(file_path, 'r', encoding='utf-8') as file:
            # Try to detect delimiter
            sample = file.read(1024)
            file.seek(0)
            
            sniffer = csv.Sniffer()
            delimiter = sniffer.sniff(sample).delimiter
            
            reader = csv.reader(file, delimiter=delimiter)
            rows = list(reader)
            
            if rows:
                csv_text = self._table_to_text(rows)
                if csv_text:
                    text_parts.append(csv_text)
        
        return "\n".join(text_parts)
    
    def _parse_image(self, file_path: str) -> str:
        """Parse image using OCR"""
        if not self.ocr_enabled:
            raise Exception("OCR not enabled. Install tesseract for image text extraction.")
        
        try:
            # Open image
            image = Image.open(file_path)
            
            # Preprocess image for better OCR
            processed_image = self._preprocess_image_for_ocr(image)
            
            # Extract text
            text = pytesseract.image_to_string(processed_image)
            
            # Clean extracted text
            cleaned_text = self._clean_text(text)
            
            logger.info(f"OCR extracted {len(cleaned_text)} characters from image")
            return cleaned_text
            
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            raise Exception(f"Failed to extract text from image: {str(e)}")
    
    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """
        Preprocess image for better OCR accuracy
        """
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Resize if too small (OCR works better on larger images)
        width, height = image.size
        if width < 600 or height < 800:
            new_width = int(width * 1.5)
            new_height = int(height * 1.5)
            image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Convert to grayscale
        image = image.convert('L')
        
        # Increase contrast
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.5)
        
        return image
    
    def _dataframe_to_text(self, df: pd.DataFrame, sheet_name: str) -> str:
        """Convert pandas DataFrame to readable text"""
        if df.empty:
            return ""
        
        text_parts = []
        
        # Add headers
        headers = " | ".join(str(col) for col in df.columns)
        text_parts.append(headers)
        text_parts.append("-" * len(headers))
        
        # Add data rows
        for _, row in df.iterrows():
            row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row)
            if row_text.strip():
                text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    def _excel_worksheet_to_text(self, worksheet, sheet_name: str) -> str:
        """Convert Excel worksheet to text"""
        text_parts = []
        
        # Find used range
        max_row = worksheet.max_row
        max_col = worksheet.max_column
        
        if max_row == 1 and max_col == 1:
            return ""
        
        # Extract data
        data = []
        for row in worksheet.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col, values_only=True):
            row_data = [str(val) if val is not None else "" for val in row]
            if any(row_data):
                data.append(row_data)
        
        return self._table_to_text(data) if data else ""
    
    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """Convert table data to readable format"""
        if not table_data:
            return ""
        
        text_rows = []
        for row in table_data:
            # Skip completely empty rows
            if any(cell.strip() for cell in row):
                row_text = " | ".join(cell.strip() for cell in row)
                text_rows.append(row_text)
        
        return "\n".join(text_rows) if text_rows else ""
    
    def _docx_table_to_text(self, table) -> str:
        """Convert DOCX table to text"""
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        
        return self._table_to_text(table_data)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove null characters
        text = text.replace('\x00', '')
        
        # Fix spacing around punctuation
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_document_info(self, file_path: str) -> Dict[str, Union[str, int, bool]]:
        """Get document metadata without full parsing"""
        info = {
            'file_path': file_path,
            'file_size': 0,
            'page_count': 0,
            'format': None,
            'parsing_success': False
        }
        
        try:
            info['file_size'] = os.path.getsize(file_path)
            
            _, ext = os.path.splitext(file_path)
            info['format'] = ext.lower()
            
            # Format-specific info
            if info['format'] == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    info['page_count'] = len(pdf.pages)
            
            elif info['format'] == '.docx':
                doc = Document(file_path)
                info['page_count'] = len(doc.paragraphs)  # Approximate
            
            elif info['format'] == '.xlsx':
                wb = load_workbook(file_path, read_only=True)
                info['page_count'] = len(wb.sheetnames)
            
            elif info['format'] == '.csv':
                with open(file_path, 'r') as f:
                    info['page_count'] = sum(1 for _ in f)
            
            info['parsing_success'] = True
            
        except Exception as e:
            logger.error(f"Failed to get document info: {e}")
            info['error'] = str(e)
        
        return info